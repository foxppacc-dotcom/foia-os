#!/usr/bin/env python3
"""Migration Plan — Sprint Zero: Communication Engine Refactor"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Arial'; st.font.size = Pt(10)

def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def H3(t): doc.add_heading(t, level=3)
def P(t): doc.add_paragraph(t)
def L(t): doc.add_paragraph(t, style='List Bullet')
def C(t, rows):
    n = len(rows[0]) if rows else 1
    tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.style = 'Light Grid Accent 1'
    for ci, h in enumerate(t):
        run = tbl.rows[0].cells[ci].paragraphs[0].add_run(h); run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri+1].cells[ci].text = str(val)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Sprint Zero — Communication Engine Refactor'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(212, 168, 67)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Migration Plan & Dependency Audit'); r.font.size = Pt(13); r.font.color.rgb = RGBColor(150, 150, 160)
doc.add_paragraph()

# ════════════════════════════════════════
H1('Phase 1: Current State Audit')
# ════════════════════════════════════════

H2('1.1 Route Files — Communication Responsibility')
C(['Route File', 'Responsibility', 'Duplicated With', 'Status'],
  [
      ['documentCenter.js', 'Compose, IMAP Poll, IMAP Diagnostic, fix-credentials, Email Account CRUD', 'email.js (send)', '❌ DUPLICATE'],
      ['email.js', 'Send, Email Account CRUD', 'documentCenter.js (compose, imap)', '❌ DUPLICATE'],
      ['communications.js', 'GET case communications', 'inbox endpoint', '⚠️ PARTIAL'],
      ['mailPoller.js', 'IMAP poll + match + store', 'imapService.js (diagnose)', '❌ DUPLICATE'],
      ['imapService.js', 'IMAP diagnose + connectivity + compare', 'mailPoller.js (poll)', '❌ DUPLICATE'],
      ['emailService.js', 'SMTP send helper', 'documentCenter.js (inline SMTP)', '❌ DUPLICATE'],
  ])

H2('1.2 Service Files')
C(['Service', 'Responsibility', 'Issues'],
  [
      ['emailService.js', 'nodemailer transport + send', 'Used by email.js only. Not used by documentCenter.js (inline SMTP)'],
      ['mailPoller.js', 'IMAP poll + download + parse + match + store', 'Broken import (fix applied). Masks errors. SQLite-dependent.'],
      ['imapService.js', 'IMAP diagnose + connectivity + compare + fix-credentials', 'Well-instrumented, proper error handling. pollAccount() stubbed.'],
      ['crypto.js', 'AES-256-GCM encrypt/decrypt', 'Used by email.js + documentCenter.js'],
  ])

H2('1.3 SMTP Implementations — COUNT: 3')
L('documentCenter.js:148-205 — inline nodemailer in compose (primary)')
L('email.js:160-198 — emailService.sendEmail() (secondary)')
L('emailService.js:60-95 — nodemailer transport builder (shared only by email.js)')
P('▸ Analysis: The compose endpoint bypasses emailService.js entirely, creating its own nodemailer transport. This is the PRIMARY source of SMTP send but is NOT shared.')

H2('1.4 IMAP Implementations — COUNT: 2')
L('mailPoller.js — full poll cycle (broken import, error masking, SQLite)')
L('imapService.js — diagnose + connectivity + compare (well-instrumented, proper errors)')
P('▸ Analysis: imapService.js is the superior implementation but pollAccount() is stubbed. mailPoller.js has the poll logic but is poorly structured.')

H2('1.5 Authentication Implementations — COUNT: 2')
L('documentCenter.js (compose): decrypt(account.smtp_pass) inline')
L('emailService.js: decrypt(account.smtp_pass) in Service')
P('▸ Both use the same crypto.js, but the transport creation is duplicated.')

H2('1.6 Auto-Matching Implementations — COUNT: 1 (insufficient)')
L('mailPoller.js:68-105 — tries Message-ID, In-Reply-To, References, Agency, Contact, Case Number')
L('No matching endpoint exists for manual re-match')
P('▸ Matching logic exists only inside the poller. No API to re-match existing records.')

H2('1.7 Timeline Implementations — COUNT: 2')
L('documentCenter.js:245-249 — inserts into case_comments on compose')
L('mailPoller.js:128-135 — tries to insert into case_activities or case_comments')
P('▸ Second implementation uses wrong table (case_activities) and fails silently (.catch(() => {}))')

H2('1.8 Conversation Architecture — MISSING')
L('No conversations table exists')
L('Messages stored directly in communications with case_id FK')
L('No conversation_id column')
L('Thread grouping is by thread_id only (stored inconsistently)')
P('▸ CRITICAL GAP: The conversation abstraction does not exist in the database.')

H2('1.9 Duplicate Routes')
C(['Endpoint', 'File', 'Duplicate At', 'Conflict'],
  [
      ['POST /api/cases/:id/compose', 'documentCenter.js', 'POST /api/send (email.js)', 'Two ways to send email'],
      ['GET /api/email-accounts', 'documentCenter.js', 'GET /api/email-accounts (email.js)', 'TWO handlers for same path!'],
      ['POST /api/imap/poll (old)', 'documentCenter.js', 'Deprecated (uses mailPoller)', 'Should use imapService'],
      ['PUT /api/email-accounts/:id', 'email.js', 'Same in documentCenter.js', 'Possible route collision'],
  ])

H1('Phase 2: Target Architecture')

H2('2.1 New Dependency Graph')
P('CommunicationService (SINGLETON)')
L('├─ SMTP Transport (nodemailer, single instance)')
L('├─ IMAP Transport (imapflow, single instance)')
L('├─ Matching Engine (Message-ID → In-Reply-To → References → Existing → Contact → Agency → Case)')
L('├─ Conversation Engine (create, update, status, archive)')
L('├─ Timeline Engine (case_comments inserts)')
L('├─ Notification Engine (notifications table)')
L('└─ Attachment Engine (Supabase Storage)')

P('API Routes → CommunicationService only')
L('├─ /api/communications/send')
L('├─ /api/communications/reply')
L('├─ /api/communications/fetch (IMAP poll)')
L('├─ /api/communications/archive')
L('├─ /api/communications/diagnose')
L('├─ /api/conversations (list)')
L('├─ /api/conversations/:id (detail)')
L('└─ /api/email-accounts (CRUD — single handler)')

H2('2.2 Database Changes')
C(['Table', 'Column', 'Type', 'Purpose'],
  [
      ['communications', 'conversation_id', 'INTEGER FK→conversations', '(NEW) Link to conversation'],
      ['communications', 'message_id', 'TEXT', 'SMTP Message-ID (FIXED)'],
      ['communications', 'thread_id', 'TEXT', 'Thread identifier (FIXED)'],
      ['communications', 'in_reply_to', 'TEXT', 'Parent Message-ID'],
      ['communications', 'references', 'TEXT', 'Reference chain'],
      ['communications', 'is_archived', 'BOOLEAN', 'Archive flag (NEW)'],
      ['communications', 'email_account_id', 'INTEGER FK', 'Source account (NEW)'],
      ['communications', 'matched_at', 'TIMESTAMP', 'When auto-matched (NEW)'],
      ['communications', 'match_method', 'TEXT', 'How matched (NEW)'],
      ['conversations', 'id', 'SERIAL PK', '(NEW TABLE)'],
      ['conversations', 'investigation_id', 'INTEGER FK', '(NEW)'],
      ['conversations', 'agency_id', 'INTEGER FK', '(NEW)'],
      ['conversations', 'contact_id', 'INTEGER FK', '(NEW)'],
      ['conversations', 'thread_id', 'TEXT', '(NEW) Root thread ID'],
      ['conversations', 'status', 'TEXT', '(NEW) open/waiting/closed'],
      ['conversations', 'unread_count', 'INTEGER', '(NEW)'],
      ['conversations', 'last_activity', 'TIMESTAMP', '(NEW)'],
      ['notifications', 'id', 'SERIAL PK', '(NEW TABLE)'],
      ['notifications', 'investigation_id', 'INTEGER FK', '(NEW)'],
      ['notifications', 'message_id', 'INTEGER FK→communications', '(NEW)'],
      ['notifications', 'type', 'TEXT', '(NEW) email_received/reply'],
      ['notifications', 'is_read', 'BOOLEAN', '(NEW)'],
  ])

H2('2.3 API Changes')
C(['Old Endpoint', 'New Endpoint', 'Notes'],
  [
      ['POST /cases/:id/compose', 'POST /communications/send', 'Unified send'],
      ['POST /api/send', 'POST /communications/send', 'Same handler'],
      ['POST /api/imap/poll', 'POST /communications/fetch', 'Renamed, same logic'],
      ['GET /api/inbox', 'GET /communications', 'Unified inbox'],
      ['GET /api/cases/:id/communications', 'GET /communications?case_id=:id', 'Filtered view'],
      ['N/A', 'POST /communications/reply', 'New — reply within conversation'],
      ['N/A', 'PATCH /conversations/:id/status', 'Status update'],
      ['N/A', 'POST /communications/match/:id', 'Manual re-match'],
      ['N/A', 'POST /conversations/:id/archive', 'Archive'],
  ])

H1('Phase 3: Migration Strategy')

H2('3.1 Order of Operations')
P('Phase A — Foundation (no breaking changes):')
L('A1. Create conversations table in Supabase')
L('A2. Add conversation_id + is_archived + email_account_id + match_method columns to communications')
L('A3. Create CommunicationService.js with unified SMTP + IMAP + Matching')
L('A4. Create notifications table')
L('A5. Run data migration: create conversations from existing thread_id groups')

P('Phase B — Route Consolidation:')
L('B1. Create /api/communications/* routes (single router)')
L('B2. Create /api/conversations/* routes')
L('B3. Move email account CRUD to single handler (remove email.js duplicate)')
L('B4. Keep old endpoints working as redirects or remove after testing')

P('Phase C — Frontend Integration:')
L('C1. Create CommunicationRepository.js')
L('C2. Update Global Inbox to use repository')
L('C3. Update CommunicationCenter to use repository')
L('C4. Update Agency page to use repository')
L('C5. Verify all views show identical data')

P('Phase D — Cleanup:')
L('D1. Remove mailPoller.js (fully replaced by imapService + CommunicationService)')
L('D2. Remove email.js duplicate routes')
L('D3. Remove documentCenter.js inline SMTP (now in CommunicationService)')
L('D4. Consolidate emailService.js into CommunicationService')

H2('3.2 Risk Assessment')
C(['Risk', 'Severity', 'Mitigation'],
  [
      ['Conversations migration: thread_id collisions', 'Medium', 'Backup before migration, dry-run'],
      ['Old routes removed before frontend updated', 'High', 'Keep old routes, add new ones alongside'],
      ['Data loss during column additions', 'Low', 'ALTER TABLE ADD COLUMN, not destructive'],
      ['Frontend timing: stale API calls', 'Medium', 'Phase C after B, test on staging'],
      ['IMAP poller downtime during swap', 'Low', 'Swap happens in one deploy'],
  ])

H2('3.3 Production Readiness Score (Current)')
C(['Subsystem', 'Score', 'Notes'],
  [
      ['SMTP', '7/10', '3 implementations, but each works. message_id now stored.'],
      ['IMAP', '6/10', 'Connection works. Poll logic in old mailPoller. imapService is better.'],
      ['Auto Matching', '2/10', 'Only runs inside poller. No API. 0/20 matched historically.'],
      ['Conversations', '0/10', 'Table does not exist. No conversation_id.'],
      ['Timeline', '4/10', 'Works for outbound. Broken for inbound (wrong table).'],
      ['Global Inbox', '6/10', 'Works but schema inconsistent (msg_id vs message_id).'],
      ['Communication Center', '5/10', 'Works for outbound, missing inbound matching.'],
      ['Agency History', '3/10', 'No agency_id on inbound records.'],
      ['Duplicate Prevention', '3/10', 'Depends on message_id (was NULL). Now fixed.'],
      ['Notifications', '0/10', 'No notification table or engine.'],
      ['Error Handling', '3/10', 'Catch(){}, return [], success:true with hidden errors.'],
      ['Code Duplication', '2/10', '3 SMTP impls, 2 IMAP impls, 2 email account CRUD endpoints.'],
      ['Overall', '38/120 (32%)', 'Production-critical features missing or duplicated.'],
  ])

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of Migration Plan —'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(150, 150, 160)

out = os.path.expanduser('~/Desktop/التطوير/99_MigrationPlan.docx')
doc.save(out)
print('DOCX:', out)
