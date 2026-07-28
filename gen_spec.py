#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──
st = doc.styles['Normal']; st.font.name = 'Arial'; st.font.size = Pt(10)
for level in range(1, 7):
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Arial'

def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def H3(t): doc.add_heading(t, level=3)
def P(t): doc.add_paragraph(t)
def L(t): p = doc.add_paragraph(t, style='List Bullet'); p.paragraph_format.space_after = Pt(2)
def NOTE(t): p = doc.add_paragraph(); r = p.add_run(f'⚠️ {t}'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(200, 150, 50)

def T(h, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(h)); t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(h): t.rows[0].cells[i].text = hd
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row): t.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

# ═══════════════════════════════ TITLE ═══════════════════════════════
title_p = doc.add_paragraph(); title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('FOIA OS v2 — Enterprise Communication Architecture Specification'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(212, 168, 67)
sub_p = doc.add_paragraph(); sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run('Sprint 18 · Phase 0 · Architecture Review · Version 1.0'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(150, 150, 160)
doc.add_paragraph()
P('Author: Hermes Agent / Amr'); P('Date: July 2026'); P('Status: Draft — Awaiting Approval')
doc.add_paragraph()

# ═══════════════════════════ 1. SOURCE OF TRUTH ═══════════════════════════
H1('1. Source of Truth')
P('The Global Mailbox is the ONE authoritative source of truth for all communications.')
P('Every email — incoming, outgoing, linked, unlinked — lives in the communications table. Investigations, Agency pages, Timeline, and Contact history are all FILTERED VIEWS of the same global dataset.')
P('There is only ONE record per email. No duplication. No per-investigation message copies.')
P('Rule: A message is created once in communications. All other contexts (Investigation, Agency, Timeline, Mailbox folders) reference the same row via case_id, agency_id, or conversation_id.')

NOTE('Current violation: No conversations table exists. Messages are stored directly in communications with a thread_id but no conversation_id FK. This creates the risk of orphaned threads.')

# ═══════════════════════════ 2. ENTITIES ═══════════════════════════
H1('2. Entity Definitions')

H2('2.1 Email Account')
T(['Property', 'Value'],
  [
      ('Purpose', 'SMTP/IMAP configuration for sending and receiving email through a single mailbox'),
      ('Owner', 'Organization (system admin)'),
      ('Table', 'email_accounts'),
      ('Relationships', 'Has many communications (via email_account_id). Assigned to requests/cases via assigned_email_account_id.'),
      ('Lifecycle', 'Created by admin → verified (SMTP/IMAP test) → active → disabled/deleted'),
      ('Key fields', 'email, smtp_host/port/user/pass (encrypted), imap_host/port/user/pass (encrypted), daily_limit, is_active, provider'),
  ])

H2('2.2 Agency')
T(['Property', 'Value'],
  [
      ('Purpose', 'External organization that FOIA requests are sent to'),
      ('Owner', 'Organization (admin/manager)'),
      ('Table', 'agencies'),
      ('Relationships', 'Has many contacts (via agencies.notes._contacts JSON or agency_contacts table). Assigned to investigations via requests.'),
      ('Lifecycle', 'Created → activated → linked to investigations → communication history accumulates'),
      ('Key fields', 'name_en, name_ar, email, phone, state, city, type, is_active, default_email_account_id, notes'),
  ])
NOTE('Current violation: Agency contacts stored in agencies.notes JSON field instead of a proper agency_contacts table. Needs migration.')

H2('2.3 Agency Contact')
T(['Property', 'Value'],
  [
      ('Purpose', 'Specific person within an Agency who handles communications'),
      ('Owner', 'Agency'),
      ('Table', '(should be) agency_contacts — currently stored in agencies.notes._contacts JSON'),
      ('Relationships', 'Belongs to Agency. Referenced by conversations (via contact_id).'),
      ('Lifecycle', 'Created → active → designated as default for an Agency → deactivated'),
      ('Key fields', 'id, agency_id, name, title, email, phone, mobile, notes, is_active, is_default'),
  ])

H2('2.4 Investigation')
T(['Property', 'Value'],
  [
      ('Purpose', 'Primary operational unit — a FOIA case that requires communication with agencies'),
      ('Owner', 'Investigation (itself)'),
      ('Table', 'cases'),
      ('Relationships', 'Has many requests (agency assignments). Has many conversations (via investigation_id). Has many communications (filtered view). Has timeline (case_comments).'),
      ('Lifecycle', 'Created → assigned → agencies added → communications sent/received → evidence collected → closed'),
      ('Key fields', 'id, title, description, status, priority, owner_id, created_at'),
  ])

H2('2.5 Conversation')
T(['Property', 'Value'],
  [
      ('Purpose', 'Central operational object that groups all messages between an Investigation and an Agency'),
      ('Owner', 'Investigation'),
      ('Table', '**(DOES NOT EXIST YET)** — proposed: conversations'),
      ('Relationships', 'Belongs to Investigation. Belongs to Agency. References a Contact. Has many Messages.'),
      ('Lifecycle', 'Created on first send → status=open → waiting for reply → closed upon resolution → archived'),
      ('Key fields', 'id, investigation_id, agency_id, contact_id, thread_id, status, waiting_since, last_activity, assigned_investigator_id, unread_count'),
      ('Statuses', 'open, waiting, closed, archived'),
  ])
NOTE('MISSING TABLE: conversations does not exist. This is the single most critical missing piece. Messages are currently stored directly in communications with no conversation grouping.')

H2('2.6 Message')
T(['Property', 'Value'],
  [
      ('Purpose', 'A single email — sent or received. The atomic unit of communication.'),
      ('Owner', 'Conversation (or Global Mailbox if unlinked)'),
      ('Table', 'communications'),
      ('Relationships', 'Belongs to a Conversation (via conversation_id). Optionally linked to Investigation (case_id) and Agency (agency_id). Has attachments.'),
      ('Lifecycle', 'Created when sent via SMTP or received via IMAP → stored once → visible in all views'),
      ('Key fields', 'id, conversation_id, case_id, agency_id, email_account_id, thread_id, message_id, in_reply_to, references, direction (outbound/inbound/draft/system), sender, recipient, cc, bcc, subject, body, html_body, attachments, is_read, is_archived, created_at'),
      ('Directions', 'outbound, inbound, draft, system'),
  ])
NOTE('Current violation: conversation_id column does not exist on communications. Messages use thread_id for grouping.')

H2('2.7 Attachment')
T(['Property', 'Value'],
  [
      ('Purpose', 'File sent with or received alongside an email'),
      ('Owner', 'Message'),
      ('Table', '(stored in communications.file_paths JSON array or documents table)'),
      ('Relationships', 'Belongs to one Message. Can be promoted to Evidence or Document.'),
      ('Lifecycle', 'Downloaded/uploaded → stored → optionally linked to Evidence → retained'),
  ])

H2('2.8 Timeline Event')
T(['Property', 'Value'],
  [
      ('Purpose', 'Activity log entry showing communication events inside an Investigation'),
      ('Owner', 'Investigation'),
      ('Table', 'case_comments (with activity icons 📧/📄)'),
      ('Relationships', 'Belongs to Investigation (case_id). References Message (via content or target_id).'),
      ('Lifecycle', 'Created when message sent/received/read → displayed in Investigation timeline'),
  ])
NOTE('Current implementation: Uses case_comments with icon markers (📧/📄). Marginally acceptable but should evolve to a dedicated case_activities table.')

H2('2.9 Notification')
T(['Property', 'Value'],
  [
      ('Purpose', 'Alert the assigned investigator of new communication events'),
      ('Owner', 'Investigator / User'),
      ('Table', 'notifications (or realtime push)'),
      ('Relationships', 'Belongs to User. References Message and Conversation.'),
      ('Lifecycle', 'Created on incoming message → displayed as badge → dismissed → cleared'),
  ])

# ═══════════════════════════ 3. FLOWS ═══════════════════════════
doc.add_page_break()
H1('3. Data Flows')

H2('3.1 Outgoing Email Flow')
P('Step-by-step for sending an email from an Investigation:')
L('1. Investigator opens Investigation → Communication Center')
L('2. Selects Agency (dropdown populated from investigation agencies)')
L('3. System auto-selects default Contact for that Agency')
L('4. System auto-fills recipient = contact.email')
L('5. Investigator can change Contact or manually type recipient')
L('6. System auto-selects assigned Email Account (from request.communication-config)')
L('7. Investigator composes subject + body + optional CC/BCC/attachments')
L('8. Clicks Send')
L('9. Backend:')
L('   a. Decrypts smtp_pass from the email account')
L('   b. Creates nodemailer transporter')
L('   c. Sends via SMTP')
L('   d. On success, creates Message record in communications table')
L('   e. Message gets: message_id (from SMTP), thread_id (from References), direction=outbound')
L('   f. Finds or creates Conversation:')
L('      - If thread_id exists in an open conversation → use it')
L('      - Otherwise → create new Conversation linked to Investigation + Agency + Contact')
L('   g. Links Message to Conversation (conversation_id)')
L('   h. Creates Timeline event in case_comments (📧 subject)')
L('   i. Message is now visible in:')
L('      - Global Mailbox (all communications)')
L('      - Investigation Communication Center (filter by case_id)')
L('      - Agency communication history (filter by agency_id)')
L('      - Conversation thread (filter by conversation_id)')
L('      - Timeline (from case_comments)')
L('10. Frontend refreshes thread list')

H2('3.2 Incoming Email Flow')
P('Automatic process triggered by IMAP poll (scheduled cron job or manual):')
L('1. Cron triggers POST /api/imap/poll for each active email account')
L('2. For each account:')
L('   a. Connect via IMAP (using decrypted credentials)')
L('   b. Search for unseen messages in INBOX')
L('   c. For each unseen message:')
L('      i. Download headers + body + attachments')
L('      ii. Parse Message-ID, In-Reply-To, References, From, To, Subject')
L('      iii. Check messageId cache for duplicate prevention')
L('      iv. If duplicate → skip')
L('      v. Create Message record in communications table (direction=inbound)')
L('      vi. Auto-match to Conversation:')
L('         - Priority 1: Message-ID matches sent message → find that message\'s conversation')
L('         - Priority 2: In-Reply-To matches a sent message → find that conversation')
L('         - Priority 3: References matches a sent thread → find that conversation')
L('         - Priority 4: Agency Contact email matches → find matching conversation')
L('         - Priority 5: Agency email matches → find matching conversation')
L('         - Priority 6: Investigation number in Subject → find that investigation')
L('      vii. If matched → link Message to Conversation')
L('      viii. If NOT matched → leave Message unlinked (conversation_id=NULL)')
L('      ix. Update Conversation status to "waiting" if matched')
L('      x. Increment conversation.unread_count')
L('      xi. Create Timeline event (📨 reply received)')
L('      xii. Create notification for assigned investigator')
L('      xiii. Mark as unseen on server (or leave as seen)')

H2('3.3 Reply Flow (from Investigation)')
P('Investigator replies to an existing thread within the Communication Center:')
L('1. Opens existing Conversation')
L('2. Clicks Reply')
L('3. System auto-fills: To (original sender), Subject (Re: original)')
L('4. Composes body')
L('5. Clicks Send')
L('6. Backend:')
L('   a. SMTP send (same as 3.1 steps 9a-9c)')
L('   b. Creates Message with in_reply_to = original message_id')
L('   c. Links to SAME Conversation (same conversation_id)')
L('   d. Same thread_id preserved')
L('7. Message appears in same Conversation, sent folder, timeline')

H2('3.4 Incoming Reply Flow')
P('When a reply arrives via IMAP:')
L('1. IMAP poll downloads the message')
L('2. In-Reply-To matches a sent message\'s message_id')
L('3. System finds the Conversation that contains the original sent message')
L('4. Links the reply to that same Conversation')
L('5. Conversation status stays "waiting" or re-opens if was closed')
L('6. Unread count increments')
L('7. Notification sent to assigned investigator')

# ═══════════════════════════ 4. SYNCHRONIZATION ═══════════════════════════
H1('4. Synchronization Rules')
L('A message is stored ONCE in the communications table.')
L('Every view (Investigation, Agency, Timeline, Sent, Inbox) queries the SAME row.')
L('There is NO per-view duplication. No copy-on-write. No sync jobs between views.')
L('conversation_id, case_id, agency_id are the join keys that enable filtered views.')
L('If case_id is NULL → message appears only in Global Mailbox (Unlinked).')
L('If case_id is set → message appears both in Global Mailbox AND Investigation.')
L('If agency_id is set → message appears in Agency communication history.')
L('Timeline events (case_comments) are lightweight markers. The full message data lives in communications.')

# ═══════════════════════════ 5. OWNERSHIP ═══════════════════════════
H1('5. Ownership Rules')
T(['Artifact', 'Owner', 'Stored In'],
  [
      ('Thread-ID', 'Conversation', 'conversations.thread_id'),
      ('Message-ID', 'Message (SMTP server)', 'communications.message_id'),
      ('Conversation-ID', 'Investigation', 'conversations.id → communications.conversation_id'),
      ('Attachments', 'Message', 'communications.file_paths (JSON) or documents table'),
      ('Read status', 'User (per message)', 'communications.is_read'),
      ('Archive status', 'User (per message)', 'communications.is_archived'),
      ('Waiting status', 'Conversation', 'conversations.status + waiting_since'),
      ('Linked Investigation', 'Message (via case_id)', 'communications.case_id'),
      ('Linked Agency', 'Message (via agency_id)', 'communications.agency_id'),
      ('Linked Contact', 'Message (via conversation.contact_id)', 'conversations.contact_id'),
  ])

# ═══════════════════════════ 6. FAILURE HANDLING ═══════════════════════════
H1('6. Failure Handling')
H2('SMTP Failure')
L('On send failure: Message is NOT stored. User sees error message.')
L('Retry: User clicks Send again. No orphan records.')
L('Partial failure: If SMTP accepts but nodemailer errors on response, store message with status="send_pending" and a retry job handles it.')

H2('IMAP Failure')
L('Connection error: Log error. Skip account. Try next account. Retry on next poll cycle.')
L('Auth error: Mark account as "imap_error". Notify admin.')
L('Rate limit: Exponential backoff. Max retry 3 times, then skip until next cycle.')
L('Partial sync: If download fails for one message, skip and continue with remaining. Re-attempt on next poll.')

H2('Duplicate Email')
L('Prevention: Check messageId cache (Set of recently seen message_ids) before storing.')
L('The cache stores message_ids seen in the last N poll cycles (rolling window).')
L('If in_reply_to matches a stored message_id, and the message content is identical → skip.')

H2('Unknown Sender')
L('After all 7 matching priorities fail → message remains unlinked.')
L('Unlinked messages appear in Global Mailbox → user can manually link or create investigation.')
L('No email is ever lost. Every message is stored before matching is attempted.')

H2('Deleted Investigation / Agency')
L('If investigation is deleted: messages remain in Global Mailbox (case_id = NULL). No data loss.')
L('If agency is deleted: messages remain in Global Mailbox (agency_id = NULL). No data loss.')
L('Conversations are soft-deleted (status=archived). Messages are never cascade-deleted.')

H2('Changed Contact')
L('Old messages retain original contact reference. New messages use new contact.')
L('Conversation.contact_id can be updated to reflect the new primary contact.')

H2('Network Interruption')
L('SMTP: timeout after 30s. Retry is manual (user clicks Send again).')
L('IMAP: timeout after 30s per folder. Errors logged. Retries on next cycle.')
L('API: All backend endpoints return JSON errors. Frontend shows user-friendly message.')

# ═══════════════════════════ 7. GAP ANALYSIS ═══════════════════════════
doc.add_page_break()
H1('7. Gap Analysis: Specification vs Current Implementation')

H2('7.1 Missing Tables')
T(['Table', 'Required By', 'Current State'],
  [
      ('conversations', 'Architecture §2.5', 'DOES NOT EXIST — messages stored directly in communications'),
      ('agency_contacts', 'Architecture §2.3', 'DOES NOT EXIST — contacts stored in agencies.notes JSON'),
  ])

H2('7.2 Missing Columns on Existing Tables')
T(['Table', 'Missing Column', 'Why Needed'],
  [
      ('communications', 'conversation_id', 'FK to conversations table — enables threaded grouping'),
      ('communications', 'html_body', 'Store rich HTML email body'),
      ('communications', 'cc', 'Carbon-copy recipients'),
      ('communications', 'bcc', 'Blind carbon-copy recipients'),
      ('communications', 'is_read', 'Read/unread tracking'),
      ('communications', 'is_archived', 'Archive flag for mailbox views'),
      ('communications', 'in_reply_to', 'Reply chain tracking'),
      ('communications', 'references', 'Thread reference chain'),
  ])

H2('7.3 Duplicate Backend Routes')
T(['Route', 'File', 'Duplicate File'],
  [
      ('/email-accounts', 'email.js', 'documentCenter.js (same route, registered later → wins)'),
      ('/inbox', 'email.js', 'documentCenter.js (same route, registered later → wins)'),
      ('/send', 'email.js', '(no direct duplicate, but /cases/:id/compose in documentCenter.js does the same thing)'),
  ])

H2('7.4 Duplicate SMTP Send Logic')
T(['Location', 'How It Sends', 'Uses emailService?'],
  [
      ('email.js → /send', 'emailService.sendEmail() with decrypt', 'Yes'),
      ('documentCenter.js → /cases/:id/compose', 'Inline nodemailer with decrypt', 'No — duplicated logic'),
      ('emailProduction.js → /email/test-compose', 'Inline nodemailer with decrypt', 'No — duplicated logic'),
  ])

H2('7.5 Frontend Inconsistencies')
T(['Issue', 'Details'],
  [
      ('API URL normalization', '8 files had no /api normalization. Fixed with getApiBase(). But some files still have inline fetch() with hardcoded response formats.'),
      ('Response format parsing', 'EmailAccounts page expects {accounts:...}, AgenciesTab expects {accounts:...}, backend returns {data:...}. Fixed in CommunicationCenter and AgenciesTab. Others may still be broken.'),
      ('Token key', 'localStorage "token" vs "foia_token" — fixed in EmailAccounts, CommunicationCenter, Inbox. SettingsTabs.jsx, OrganizationsHub.jsx, TeamTab.jsx, AgenciesTab.jsx still use "token".'),
      ('No shared fetch helper', 'Many components define their own hdrs(), tok(), API inline instead of importing from api.js.'),
  ])

H2('7.6 Backend Inconsistencies')
T(['Issue', 'Details'],
  [
      ('IMAP poll not scheduled', 'mailPoller.js built but no cron job running it automatically'),
      ('agency_contacts table missing', 'All references to agency_contacts will fail with "table not found"'),
      ('case_activities table unreliable', 'Migrated to case_comments but some endpoints still reference case_activities'),
      ('No Conversation endpoints', 'No POST/GET/PUT /conversations routes'),
      ('No matching engine endpoint', 'Auto-match logic is in mailPoller.js but not exposed as standalone API'),
  ])

H2('7.7 Remaining localStorage "token" Bugs')
T(['File', 'Line', 'Bug'],
  [
      ('AgenciesTab.jsx', '16', 'localStorage.getItem("token")'),
      ('CaseHeader.jsx', '12', 'localStorage.getItem("token")'),
      ('DocumentsTab.jsx', '9', 'localStorage.getItem("token")'),
      ('TeamTab.jsx', '12', 'localStorage.getItem("token")'),
      ('SettingsTabs.jsx', '10', 'localStorage.getItem("token")'),
      ('OrganizationHub.jsx', '10', 'localStorage.getItem("token")'),
      ('WorkloadDashboard.jsx', '~15', 'localStorage.getItem("token")'),
  ])

# ═══════════════════════════ 8. PRIORITY ═══════════════════════════
H1('8. Recommended Implementation Order')
P('Based on the gap analysis, the logical order of work:')
L('P0 — Add conversations table + conversation_id FK on communications (gap §7.1, §7.2)')
L('P0 — Consolidate duplicate backend routes (gap §7.3)')
L('P0 — Consolidate SMTP send logic into emailService (gap §7.4)')
L('P0 — Create agency_contacts table + migrate data from notes JSON (gap §7.1)')
L('P0 — Fix localStorage "token" → "foia_token" in remaining 7 files (gap §7.7)')
L('P1 — Add conversation endpoints (CRUD + matching)')
L('P1 — Schedule IMAP polling cron job')
L('P1 — Add missing columns to communications (html_body, cc, bcc, is_read, is_archived, in_reply_to, references)')
L('P1 — Build ConversationView UI component')
L('P2 — Threaded conversation display')
L('P2 — Notifications for incoming mail')
L('P2 — Attachments as Evidence')
L('P3 — Agency+Contact communication history pages')
L('P3 — Rich text compose + CC/BCC')
L('P3 — Message-ID cache + duplicate prevention production hardening')

# ═══════════════════════════ FOOTER ═══════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of Architecture Specification —'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(150, 150, 160)

out = os.path.expanduser('~/Desktop/التطوير/00_CommArchitectureSpec.docx')
doc.save(out)
print('DOCX:', out)
